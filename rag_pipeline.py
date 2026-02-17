#!/usr/bin/env python3
"""
RAG Pipeline for LAAA BOV Chatbot
==================================
Document parsing, chunking, embedding, vector upload, and chat widget generation.
Used by build_bov_v2.py to create a property-specific AI chatbot for each BOV website.

Usage:
    This module is imported by build_bov_v2.py. It is NOT run directly.
    All API keys are loaded from .env at import time.
"""

import os
import re
import json
import hashlib
from dataclasses import dataclass, field, asdict
from typing import Optional
from pathlib import Path

# ---------------------------------------------------------------------------
# Lazy imports: heavy libraries are only imported when their functions are called.
# This prevents build_bov_v2.py from crashing if optional deps aren't installed
# when ENABLE_CHATBOT = False.
# ---------------------------------------------------------------------------


# ============================================================
# DATA CLASSES
# ============================================================

@dataclass
class Document:
    """A parsed document segment with source metadata."""
    text: str
    source: str          # e.g., "420-428 W Stocker - Pricing Model.pdf"
    page: str = ""       # e.g., "Page 3" or "Sheet: Rent Roll" or "Section: Financials"
    doc_type: str = ""   # e.g., "pdf", "xlsx", "website", "txt", "md", "docx"


@dataclass
class Chunk:
    """A chunked piece of text ready for embedding."""
    text: str
    source: str
    page: str = ""
    doc_type: str = ""
    chunk_id: int = 0
    is_tabular: bool = False  # True for financial tables, spreadsheets


# ============================================================
# DOCUMENT PARSERS
# ============================================================

def parse_pdf(filepath: str) -> list[Document]:
    """
    Extract text from a PDF file page-by-page.
    Uses PyMuPDF (fitz). Falls back gracefully if pages are scanned images
    (returns empty text for those pages with a warning).
    """
    import fitz  # pymupdf

    docs = []
    filename = os.path.basename(filepath)

    try:
        pdf = fitz.open(filepath)
        for page_num in range(len(pdf)):
            page = pdf[page_num]
            text = page.get_text("text")

            # Skip pages with very little text (likely scanned images)
            if text and len(text.strip()) > 30:
                docs.append(Document(
                    text=text.strip(),
                    source=filename,
                    page=f"Page {page_num + 1}",
                    doc_type="pdf"
                ))
            elif len(text.strip()) <= 30:
                # Try to extract from tables/structured content
                blocks = page.get_text("blocks")
                block_text = "\n".join(
                    b[4].strip() for b in blocks if b[6] == 0 and b[4].strip()
                )
                if block_text and len(block_text.strip()) > 30:
                    docs.append(Document(
                        text=block_text.strip(),
                        source=filename,
                        page=f"Page {page_num + 1}",
                        doc_type="pdf"
                    ))
        pdf.close()
    except Exception as e:
        print(f"  WARNING: Could not parse PDF {filename}: {e}")

    return docs


def parse_docx(filepath: str) -> list[Document]:
    """Extract text from a Word document, including tables."""
    import docx

    docs = []
    filename = os.path.basename(filepath)

    try:
        doc = docx.Document(filepath)

        # Extract paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())

        if paragraphs:
            docs.append(Document(
                text="\n\n".join(paragraphs),
                source=filename,
                page="Full Document",
                doc_type="docx"
            ))

        # Extract tables as separate documents (they're often financial data)
        for i, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                table_text = "\n".join(rows)
                if len(table_text.strip()) > 20:
                    docs.append(Document(
                        text=table_text,
                        source=filename,
                        page=f"Table {i + 1}",
                        doc_type="docx"
                    ))
    except Exception as e:
        print(f"  WARNING: Could not parse DOCX {filename}: {e}")

    return docs


def parse_spreadsheet(filepath: str) -> list[Document]:
    """
    Extract data from Excel files (.xlsx, .xlsm).
    Reads cached values (not macro-computed). Each sheet becomes a Document.
    """
    import openpyxl

    docs = []
    filename = os.path.basename(filepath)

    try:
        # data_only=True reads cached formula results instead of formulas
        wb = openpyxl.load_workbook(filepath, data_only=True, read_only=True)

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                # Convert each cell to string, skip entirely empty rows
                cells = []
                for cell in row:
                    if cell is not None:
                        cells.append(str(cell))
                    else:
                        cells.append("")
                # Only include rows that have at least one non-empty cell
                if any(c.strip() for c in cells):
                    rows.append(" | ".join(cells))

            if rows:
                sheet_text = f"Sheet: {sheet_name}\n" + "\n".join(rows)
                if len(sheet_text.strip()) > 30:
                    docs.append(Document(
                        text=sheet_text,
                        source=filename,
                        page=f"Sheet: {sheet_name}",
                        doc_type="xlsx"
                    ))

        wb.close()
    except Exception as e:
        print(f"  WARNING: Could not parse spreadsheet {filename}: {e}")

    return docs


def parse_text(filepath: str) -> list[Document]:
    """Parse plain text or markdown files."""
    filename = os.path.basename(filepath)
    ext = filename.rsplit(".", 1)[-1].lower()

    try:
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            text = f.read().strip()

        if text and len(text) > 20:
            return [Document(
                text=text,
                source=filename,
                page="Full Document",
                doc_type=ext  # "txt" or "md"
            )]
    except Exception as e:
        print(f"  WARNING: Could not parse text file {filename}: {e}")

    return []


def capture_build_context(build_data: dict) -> list[Document]:
    """
    Capture structured content directly from build_bov_v2.py Python variables.
    This is MUCH better than parsing the generated HTML with BeautifulSoup,
    because we already have clean structured data.

    build_data should be a dict with keys like:
    {
        "property_name": "420-428 W Stocker St, Glendale, CA 91202",
        "list_price": 9350000,
        "units": 27,
        "sf": 22674,
        "rent_roll": [...],
        "sale_comps": [...],
        "financial_summary": "...",
        "sections": { "investment": "...", "location": "...", ... }
    }
    """
    docs = []

    # Property overview
    overview_parts = []
    if build_data.get("property_name"):
        overview_parts.append(f"Property: {build_data['property_name']}")
    if build_data.get("list_price"):
        overview_parts.append(f"List Price: ${build_data['list_price']:,.0f}")
    if build_data.get("units"):
        overview_parts.append(f"Units: {build_data['units']}")
    if build_data.get("sf"):
        overview_parts.append(f"Square Footage: {build_data['sf']:,.0f} SF")
    if overview_parts:
        docs.append(Document(
            text="\n".join(overview_parts),
            source="Website - Property Overview",
            page="Cover Page",
            doc_type="website"
        ))

    # Rent roll
    if build_data.get("rent_roll"):
        rr_lines = ["Unit | Type | SqFt | Current Rent | Market Rent"]
        for unit, utype, sqft, cur, mkt in build_data["rent_roll"]:
            rr_lines.append(f"{unit} | {utype} | {sqft} | ${cur:,} | ${mkt:,}")
        docs.append(Document(
            text="\n".join(rr_lines),
            source="Website - Rent Roll",
            page="Financials Section",
            doc_type="website"
        ))

    # Sale comps
    if build_data.get("sale_comps"):
        sc_lines = ["# | Address | Units | Price | $/Unit | $/SF | Cap | GRM | Date | Notes"]
        for c in build_data["sale_comps"]:
            cap_str = f"{c['cap']:.2f}%" if c.get("cap") else "n/a"
            grm_str = f"{c['grm']:.2f}" if c.get("grm") else "n/a"
            sc_lines.append(
                f"{c['num']} | {c['addr']} | {c['units']} | "
                f"${c['price']:,} | ${c['ppu']:,} | ${c['psf']:.0f} | "
                f"{cap_str} | {grm_str} | {c['date']} | {c.get('notes', '')}"
            )
        docs.append(Document(
            text="\n".join(sc_lines),
            source="Website - Sale Comparables",
            page="Sale Comps Section",
            doc_type="website"
        ))

    # Financial metrics
    if build_data.get("financial_summary"):
        docs.append(Document(
            text=build_data["financial_summary"],
            source="Website - Financial Summary",
            page="Financials Section",
            doc_type="website"
        ))

    # Operating statement
    if build_data.get("operating_statement"):
        docs.append(Document(
            text=build_data["operating_statement"],
            source="Website - Operating Statement",
            page="Financials Section",
            doc_type="website"
        ))

    # Free-text sections (investment overview, location, development, etc.)
    for section_name, section_text in build_data.get("sections", {}).items():
        if section_text and len(section_text.strip()) > 20:
            docs.append(Document(
                text=section_text.strip(),
                source=f"Website - {section_name}",
                page=f"{section_name} Section",
                doc_type="website"
            ))

    return docs


def parse_all_documents(docs_dir: str, build_data: Optional[dict] = None) -> list[Document]:
    """
    Master parser: iterate over all files in docs/ and dispatch to the correct parser.
    Also captures structured content from the build script if build_data is provided.

    Args:
        docs_dir: Path to the docs/ folder
        build_data: Optional dict of structured data from build_bov_v2.py
    Returns:
        List of Document objects from all sources
    """
    all_docs = []

    # Parse each file in docs/
    if os.path.isdir(docs_dir):
        for filename in sorted(os.listdir(docs_dir)):
            filepath = os.path.join(docs_dir, filename)
            if not os.path.isfile(filepath):
                continue

            ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

            # Skip images (can add Claude Vision later for Phase 2)
            if ext in ("jpeg", "jpg", "png", "gif", "webp", "bmp"):
                print(f"  Skipping image: {filename}")
                continue

            if ext == "pdf":
                parsed = parse_pdf(filepath)
            elif ext == "docx":
                parsed = parse_docx(filepath)
            elif ext in ("xlsx", "xlsm", "xls"):
                parsed = parse_spreadsheet(filepath)
            elif ext in ("txt", "md"):
                parsed = parse_text(filepath)
            else:
                print(f"  Skipping unsupported file type: {filename}")
                continue

            all_docs.extend(parsed)
            if parsed:
                print(f"  Parsed {filename}: {len(parsed)} segment(s)")

    # Capture structured data from build script variables
    if build_data:
        build_docs = capture_build_context(build_data)
        all_docs.extend(build_docs)
        print(f"  Captured {len(build_docs)} segment(s) from build script data")

    return all_docs


# ============================================================
# CHUNKING ENGINE
# ============================================================

def _count_tokens(text: str) -> int:
    """Count tokens using tiktoken (cl100k_base, used by Claude/GPT-4)."""
    import tiktoken
    enc = tiktoken.get_encoding("cl100k_base")
    return len(enc.encode(text))


def _split_sentences(text: str) -> list[str]:
    """Split text into sentences, preserving meaningful boundaries."""
    # Split on sentence-ending punctuation followed by space or newline
    sentences = re.split(r'(?<=[.!?])\s+|\n\n+', text)
    return [s.strip() for s in sentences if s.strip()]


def smart_chunk(
    documents: list[Document],
    max_tokens_narrative: int = 500,
    max_tokens_tabular: int = 900,
    overlap_tokens: int = 50
) -> list[Chunk]:
    """
    Split documents into overlapping chunks for embedding.
    Uses document-type-aware chunk sizes:
      - 500 tokens for narrative text (paragraphs, descriptions)
      - 900 tokens for tabular data (rent rolls, comps, financial tables)
    Respects sentence boundaries for narrative text.
    """
    chunks = []
    chunk_id = 0

    for doc in documents:
        # Determine if this is tabular data
        is_tabular = (
            doc.doc_type in ("xlsx", "xlsm") or
            "table" in doc.page.lower() or
            "rent roll" in doc.source.lower() or
            "operating statement" in doc.source.lower() or
            "pricing" in doc.source.lower() or
            "comps" in doc.source.lower() or
            doc.text.count("|") > 5  # pipe-delimited tables
        )

        max_tokens = max_tokens_tabular if is_tabular else max_tokens_narrative

        text = doc.text.strip()
        if not text:
            continue

        token_count = _count_tokens(text)

        # If the whole document fits in one chunk, use it as-is
        if token_count <= max_tokens:
            chunks.append(Chunk(
                text=text,
                source=doc.source,
                page=doc.page,
                doc_type=doc.doc_type,
                chunk_id=chunk_id,
                is_tabular=is_tabular
            ))
            chunk_id += 1
            continue

        # For tabular data, split on row boundaries (newlines)
        if is_tabular:
            lines = text.split("\n")
            current_chunk_lines = []
            current_tokens = 0

            for line in lines:
                line_tokens = _count_tokens(line)
                if current_tokens + line_tokens > max_tokens and current_chunk_lines:
                    chunk_text = "\n".join(current_chunk_lines)
                    chunks.append(Chunk(
                        text=chunk_text,
                        source=doc.source,
                        page=doc.page,
                        doc_type=doc.doc_type,
                        chunk_id=chunk_id,
                        is_tabular=True
                    ))
                    chunk_id += 1

                    # Overlap: keep last few lines for context
                    overlap_lines = []
                    overlap_count = 0
                    for ol in reversed(current_chunk_lines):
                        ol_tokens = _count_tokens(ol)
                        if overlap_count + ol_tokens > overlap_tokens:
                            break
                        overlap_lines.insert(0, ol)
                        overlap_count += ol_tokens

                    current_chunk_lines = overlap_lines
                    current_tokens = overlap_count

                current_chunk_lines.append(line)
                current_tokens += line_tokens

            # Final chunk
            if current_chunk_lines:
                chunk_text = "\n".join(current_chunk_lines)
                chunks.append(Chunk(
                    text=chunk_text,
                    source=doc.source,
                    page=doc.page,
                    doc_type=doc.doc_type,
                    chunk_id=chunk_id,
                    is_tabular=True
                ))
                chunk_id += 1

        else:
            # Narrative text: split on sentence boundaries
            sentences = _split_sentences(text)
            current_chunk_sentences = []
            current_tokens = 0

            for sentence in sentences:
                sent_tokens = _count_tokens(sentence)

                # If a single sentence exceeds max, force-split it
                if sent_tokens > max_tokens:
                    if current_chunk_sentences:
                        chunk_text = " ".join(current_chunk_sentences)
                        chunks.append(Chunk(
                            text=chunk_text,
                            source=doc.source,
                            page=doc.page,
                            doc_type=doc.doc_type,
                            chunk_id=chunk_id,
                            is_tabular=False
                        ))
                        chunk_id += 1
                        current_chunk_sentences = []
                        current_tokens = 0

                    # Force add the long sentence as its own chunk
                    chunks.append(Chunk(
                        text=sentence,
                        source=doc.source,
                        page=doc.page,
                        doc_type=doc.doc_type,
                        chunk_id=chunk_id,
                        is_tabular=False
                    ))
                    chunk_id += 1
                    continue

                if current_tokens + sent_tokens > max_tokens and current_chunk_sentences:
                    chunk_text = " ".join(current_chunk_sentences)
                    chunks.append(Chunk(
                        text=chunk_text,
                        source=doc.source,
                        page=doc.page,
                        doc_type=doc.doc_type,
                        chunk_id=chunk_id,
                        is_tabular=False
                    ))
                    chunk_id += 1

                    # Overlap: keep last few sentences for context
                    overlap_sents = []
                    overlap_count = 0
                    for os_ in reversed(current_chunk_sentences):
                        os_tokens = _count_tokens(os_)
                        if overlap_count + os_tokens > overlap_tokens:
                            break
                        overlap_sents.insert(0, os_)
                        overlap_count += os_tokens

                    current_chunk_sentences = overlap_sents
                    current_tokens = overlap_count

                current_chunk_sentences.append(sentence)
                current_tokens += sent_tokens

            # Final chunk
            if current_chunk_sentences:
                chunk_text = " ".join(current_chunk_sentences)
                chunks.append(Chunk(
                    text=chunk_text,
                    source=doc.source,
                    page=doc.page,
                    doc_type=doc.doc_type,
                    chunk_id=chunk_id,
                    is_tabular=False
                ))
                chunk_id += 1

    return chunks


# ============================================================
# EMBEDDING + VECTOR UPLOAD
# ============================================================

def embed_chunks(chunks: list[Chunk], namespace: str) -> list[dict]:
    """
    Generate embeddings for all chunks using Voyage AI (voyage-3, 1024 dims).
    Returns a list of Pinecone-ready vector dicts.
    Includes retry logic with exponential backoff for rate limits.

    Args:
        chunks: List of Chunk objects to embed
        namespace: BOV namespace (used as vector ID prefix)
    """
    import voyageai
    import time as _time
    from dotenv import load_dotenv
    load_dotenv()

    vo = voyageai.Client(api_key=os.getenv("VOYAGE_API_KEY"))

    vectors = []
    texts = [chunk.text for chunk in chunks]

    # Adaptive batch sizing: start small to handle rate limits, scale up if possible
    # Free tier: 3 RPM, 10K TPM. Standard tier: much higher.
    batch_size = 8  # Small batches to stay under 10K TPM on free tier
    pause_between = 21  # 21s between requests = ~2.8 RPM (under 3 RPM limit)
    all_embeddings = []
    total_batches = (len(texts) - 1) // batch_size + 1
    est_minutes = (total_batches * pause_between) / 60

    print(f"  {total_batches} batches of {batch_size} (est. {est_minutes:.0f} min with rate limit pauses)")

    for i in range(0, len(texts), batch_size):
        batch = texts[i:i + batch_size]
        batch_num = i // batch_size + 1
        print(f"  Embedding batch {batch_num}/{total_batches} ({len(batch)} chunks)...", end="", flush=True)

        # Retry with exponential backoff
        max_retries = 5
        for attempt in range(max_retries):
            try:
                result = vo.embed(batch, model="voyage-3", input_type="document")
                all_embeddings.extend(result.embeddings)
                print(" OK")
                break
            except Exception as e:
                error_msg = str(e)
                if attempt < max_retries - 1 and ("rate" in error_msg.lower() or "429" in error_msg or "limit" in error_msg.lower() or "payment" in error_msg.lower()):
                    wait = max(pause_between, (2 ** attempt) * 5)
                    print(f" rate limited, waiting {wait}s (retry {attempt + 2}/{max_retries})...", end="", flush=True)
                    _time.sleep(wait)
                else:
                    raise

        # Pause between batches to respect rate limits
        if batch_num < total_batches:
            _time.sleep(pause_between)

    for i, (embedding, chunk) in enumerate(zip(all_embeddings, chunks)):
        vectors.append({
            "id": f"{namespace}-{i:04d}",
            "values": embedding,
            "metadata": {
                "text": chunk.text[:4000],  # Pinecone metadata limit ~40KB
                "source": chunk.source,
                "page": chunk.page,
                "doc_type": chunk.doc_type,
                "is_tabular": chunk.is_tabular,
            }
        })

    return vectors


def upload_vectors(vectors: list[dict], namespace: str):
    """
    Upload vectors to Pinecone under a BOV-specific namespace.
    Clears old vectors for this namespace first (clean rebuild).
    """
    from pinecone import Pinecone
    from dotenv import load_dotenv
    load_dotenv()

    pc = Pinecone(api_key=os.getenv("PINECONE_API_KEY"))
    index = pc.Index(host=os.getenv("PINECONE_INDEX_HOST"))

    # Delete all existing vectors in this namespace (clean rebuild)
    try:
        index.delete(delete_all=True, namespace=namespace)
        print(f"  Cleared existing vectors in namespace: {namespace}")
    except Exception:
        pass  # Namespace may not exist yet

    # Upload in batches of 100
    for i in range(0, len(vectors), 100):
        batch = vectors[i:i + 100]
        index.upsert(vectors=batch, namespace=namespace)
        print(f"  Uploaded batch {i // 100 + 1}/{(len(vectors) - 1) // 100 + 1}")

    print(f"  Total vectors in namespace '{namespace}': {len(vectors)}")


# ============================================================
# CHAT WIDGET HTML/CSS/JS GENERATOR
# ============================================================

def generate_chat_widget(
    worker_url: str,
    namespace: str,
    property_name: str,
    starter_questions: list[str] = None,
    precomputed_answers: dict = None,
) -> str:
    """
    Generate a self-contained HTML/CSS/JS chat widget to be injected into index.html.

    Args:
        worker_url: URL of the Cloudflare Worker (e.g., https://laaa-chat-worker.laaa-team.workers.dev)
        namespace: Pinecone namespace for this BOV
        property_name: Display name (e.g., "420-428 W Stocker St")
        starter_questions: List of 4 suggested questions
        precomputed_answers: Optional dict mapping starter question text to pre-computed answer text
    Returns:
        HTML string with all CSS and JS inline
    """
    if starter_questions is None:
        starter_questions = [
            "What is the asking price and cap rate?",
            "Tell me about the development potential",
            "Summarize the rent roll",
            "What do the comparable sales show?"
        ]

    precomputed_json = json.dumps(precomputed_answers or {})
    starters_json = json.dumps(starter_questions)

    return f'''
<!-- ============================================================ -->
<!-- LAAA BOV AI Chat Widget -->
<!-- ============================================================ -->
<style>
/* Chat widget - hidden in print/PDF */
@media print {{ .bov-chat-widget, .bov-chat-bubble {{ display: none !important; }} }}

.bov-chat-bubble {{
    position: fixed;
    bottom: 90px;
    right: 24px;
    width: 56px;
    height: 56px;
    border-radius: 50%;
    background: #1B3A5C;
    color: #fff;
    border: none;
    cursor: pointer;
    box-shadow: 0 4px 16px rgba(0,0,0,0.25);
    z-index: 10000;
    display: flex;
    align-items: center;
    justify-content: center;
    transition: transform 0.2s, background 0.2s;
}}
.bov-chat-bubble:hover {{ transform: scale(1.08); background: #244a6e; }}
.bov-chat-bubble svg {{ width: 26px; height: 26px; fill: #fff; }}

.bov-chat-widget {{
    position: fixed;
    bottom: 90px;
    right: 24px;
    width: 390px;
    height: 540px;
    background: #fff;
    border-radius: 16px;
    box-shadow: 0 8px 40px rgba(0,0,0,0.2);
    z-index: 10001;
    display: none;
    flex-direction: column;
    overflow: hidden;
    font-family: 'Inter', -apple-system, sans-serif;
}}
.bov-chat-widget.open {{ display: flex; }}

.bov-chat-header {{
    background: #1B3A5C;
    color: #fff;
    padding: 16px 18px;
    display: flex;
    align-items: center;
    justify-content: space-between;
    flex-shrink: 0;
}}
.bov-chat-header-title {{
    font-size: 14px;
    font-weight: 600;
    line-height: 1.3;
}}
.bov-chat-header-sub {{
    font-size: 11px;
    opacity: 0.75;
    margin-top: 2px;
}}
.bov-chat-close {{
    background: none;
    border: none;
    color: #fff;
    cursor: pointer;
    font-size: 22px;
    padding: 0 0 0 12px;
    opacity: 0.8;
    line-height: 1;
}}
.bov-chat-close:hover {{ opacity: 1; }}

.bov-chat-messages {{
    flex: 1;
    overflow-y: auto;
    padding: 16px;
    display: flex;
    flex-direction: column;
    gap: 12px;
}}

.bov-chat-starters {{
    display: flex;
    flex-direction: column;
    gap: 8px;
    margin-top: 8px;
}}
.bov-chat-starter {{
    background: #f0f4f8;
    border: 1px solid #d8e1eb;
    border-radius: 10px;
    padding: 10px 14px;
    font-size: 13px;
    color: #1B3A5C;
    cursor: pointer;
    text-align: left;
    transition: background 0.15s, border-color 0.15s;
    line-height: 1.4;
}}
.bov-chat-starter:hover {{ background: #e2eaf2; border-color: #b0c4de; }}

.bov-chat-msg {{
    max-width: 88%;
    padding: 10px 14px;
    border-radius: 12px;
    font-size: 13px;
    line-height: 1.55;
    word-wrap: break-word;
}}
.bov-chat-msg.user {{
    align-self: flex-end;
    background: #1B3A5C;
    color: #fff;
    border-bottom-right-radius: 4px;
}}
.bov-chat-msg.bot {{
    align-self: flex-start;
    background: #f0f4f8;
    color: #1a1a1a;
    border-bottom-left-radius: 4px;
}}
.bov-chat-msg.bot p {{ margin: 0 0 8px 0; }}
.bov-chat-msg.bot p:last-child {{ margin-bottom: 0; }}
.bov-chat-msg.bot ul, .bov-chat-msg.bot ol {{ margin: 4px 0 8px 18px; padding: 0; }}
.bov-chat-msg.bot li {{ margin-bottom: 3px; }}
.bov-chat-msg.bot strong {{ font-weight: 600; }}
.bov-chat-msg.bot table {{ border-collapse: collapse; margin: 6px 0; width: 100%; font-size: 12px; }}
.bov-chat-msg.bot th, .bov-chat-msg.bot td {{ border: 1px solid #d0d7de; padding: 4px 8px; text-align: left; }}
.bov-chat-msg.bot th {{ background: #e8edf2; font-weight: 600; }}

.bov-chat-sources {{
    font-size: 11px;
    color: #6b7280;
    margin-top: 6px;
    cursor: pointer;
    user-select: none;
}}
.bov-chat-sources-list {{
    display: none;
    margin-top: 4px;
    padding: 6px 10px;
    background: #f9fafb;
    border-radius: 6px;
    font-size: 11px;
    line-height: 1.5;
}}
.bov-chat-sources.expanded .bov-chat-sources-list {{ display: block; }}

.bov-chat-disclaimer {{
    font-size: 10px;
    color: #9ca3af;
    margin-top: 4px;
    font-style: italic;
    line-height: 1.4;
}}

.bov-chat-typing {{
    display: flex;
    gap: 5px;
    padding: 10px 14px;
    align-self: flex-start;
    background: #f0f4f8;
    border-radius: 12px;
    border-bottom-left-radius: 4px;
}}
.bov-chat-typing span {{
    width: 7px;
    height: 7px;
    background: #94a3b8;
    border-radius: 50%;
    animation: bovTyping 1.2s ease-in-out infinite;
}}
.bov-chat-typing span:nth-child(2) {{ animation-delay: 0.2s; }}
.bov-chat-typing span:nth-child(3) {{ animation-delay: 0.4s; }}
@keyframes bovTyping {{
    0%, 60%, 100% {{ transform: translateY(0); opacity: 0.4; }}
    30% {{ transform: translateY(-6px); opacity: 1; }}
}}

.bov-chat-input-area {{
    display: flex;
    padding: 12px;
    border-top: 1px solid #e5e7eb;
    gap: 8px;
    flex-shrink: 0;
    background: #fff;
}}
.bov-chat-input {{
    flex: 1;
    border: 1px solid #d1d5db;
    border-radius: 10px;
    padding: 10px 14px;
    font-size: 13px;
    font-family: inherit;
    outline: none;
    transition: border-color 0.15s;
    resize: none;
}}
.bov-chat-input:focus {{ border-color: #1B3A5C; }}
.bov-chat-input::placeholder {{ color: #9ca3af; }}
.bov-chat-send {{
    background: #1B3A5C;
    color: #fff;
    border: none;
    border-radius: 10px;
    padding: 0 16px;
    cursor: pointer;
    font-size: 14px;
    font-weight: 600;
    transition: background 0.15s;
    flex-shrink: 0;
}}
.bov-chat-send:hover {{ background: #244a6e; }}
.bov-chat-send:disabled {{ background: #94a3b8; cursor: not-allowed; }}

.bov-chat-limit-msg {{
    text-align: center;
    padding: 12px;
    font-size: 12px;
    color: #6b7280;
    background: #f9fafb;
    border-top: 1px solid #e5e7eb;
    flex-shrink: 0;
}}
.bov-chat-limit-msg a {{ color: #1B3A5C; font-weight: 600; }}

/* Mobile responsive */
@media (max-width: 480px) {{
    .bov-chat-widget {{
        bottom: 0;
        right: 0;
        width: 100%;
        height: 100%;
        border-radius: 0;
    }}
    .bov-chat-bubble {{ bottom: 80px; right: 16px; }}
}}
</style>

<!-- Chat Bubble -->
<button class="bov-chat-bubble" id="bovChatBubble" aria-label="Ask about this property">
    <svg viewBox="0 0 24 24" xmlns="http://www.w3.org/2000/svg"><path d="M20 2H4c-1.1 0-2 .9-2 2v18l4-4h14c1.1 0 2-.9 2-2V4c0-1.1-.9-2-2-2zm0 14H6l-2 2V4h16v12z"/></svg>
</button>

<!-- Chat Panel -->
<div class="bov-chat-widget" id="bovChatWidget">
    <div class="bov-chat-header">
        <div>
            <div class="bov-chat-header-title">Ask about {property_name}</div>
            <div class="bov-chat-header-sub">AI assistant - answers from BOV materials only</div>
        </div>
        <button class="bov-chat-close" id="bovChatClose">&times;</button>
    </div>
    <div class="bov-chat-messages" id="bovChatMessages">
        <div class="bov-chat-msg bot">
            Hi! I can answer questions about the investment details, financials, comparable sales, development potential, and more for this property. What would you like to know?
        </div>
        <div class="bov-chat-starters" id="bovChatStarters"></div>
    </div>
    <div class="bov-chat-input-area" id="bovChatInputArea">
        <input type="text" class="bov-chat-input" id="bovChatInput" placeholder="Ask a question..." autocomplete="off">
        <button class="bov-chat-send" id="bovChatSend">Send</button>
    </div>
</div>

<script>
(function() {{
    'use strict';

    // --- Configuration (injected by build script) ---
    var WORKER_URL = '{worker_url}';
    var NAMESPACE = '{namespace}';
    var STARTERS = {starters_json};
    var PRECOMPUTED = {precomputed_json};
    var MAX_MESSAGES = 30;
    var DISCLAIMER = 'Information sourced from BOV materials. Verify independently before making investment decisions.';

    // --- State ---
    var history = [];
    var messageCount = 0;
    var isStreaming = false;

    // --- Elements ---
    var bubble = document.getElementById('bovChatBubble');
    var widget = document.getElementById('bovChatWidget');
    var closeBtn = document.getElementById('bovChatClose');
    var messagesEl = document.getElementById('bovChatMessages');
    var startersEl = document.getElementById('bovChatStarters');
    var inputEl = document.getElementById('bovChatInput');
    var sendBtn = document.getElementById('bovChatSend');
    var inputArea = document.getElementById('bovChatInputArea');

    // --- Render starter pills ---
    STARTERS.forEach(function(q) {{
        var pill = document.createElement('button');
        pill.className = 'bov-chat-starter';
        pill.textContent = q;
        pill.addEventListener('click', function() {{ sendMessage(q); }});
        startersEl.appendChild(pill);
    }});

    // --- Toggle chat ---
    bubble.addEventListener('click', function() {{
        widget.classList.add('open');
        bubble.style.display = 'none';
        inputEl.focus();
    }});
    closeBtn.addEventListener('click', function() {{
        widget.classList.remove('open');
        bubble.style.display = 'flex';
    }});

    // --- Send on Enter ---
    inputEl.addEventListener('keydown', function(e) {{
        if (e.key === 'Enter' && !e.shiftKey) {{
            e.preventDefault();
            sendMessage(inputEl.value);
        }}
    }});
    sendBtn.addEventListener('click', function() {{
        sendMessage(inputEl.value);
    }});

    // --- Markdown-lite renderer ---
    function renderMarkdown(text) {{
        // Escape HTML
        text = text.replace(/&/g, '&amp;').replace(/</g, '&lt;').replace(/>/g, '&gt;');
        // Bold
        text = text.replace(/\\*\\*(.+?)\\*\\*/g, '<strong>$1</strong>');
        // Inline code
        text = text.replace(/`([^`]+)`/g, '<code style="background:#e8edf2;padding:1px 4px;border-radius:3px;font-size:12px;">$1</code>');
        // Bullet lists
        text = text.replace(/^[\\-\\*] (.+)$/gm, '<li>$1</li>');
        text = text.replace(/(<li>.*<\\/li>\\n?)+/gs, '<ul>$&</ul>');
        // Numbered lists
        text = text.replace(/^\\d+[\\.\\)] (.+)$/gm, '<li>$1</li>');
        // Headers
        text = text.replace(/^### (.+)$/gm, '<strong style="font-size:14px;">$1</strong>');
        text = text.replace(/^## (.+)$/gm, '<strong style="font-size:14px;">$1</strong>');
        // Paragraphs
        text = text.replace(/\\n\\n+/g, '</p><p>');
        text = '<p>' + text + '</p>';
        text = text.replace(/<p><\\/p>/g, '');
        return text;
    }}

    // --- Add message bubble ---
    function addMessage(role, html, sources) {{
        // Remove starters after first user message
        if (role === 'user' && startersEl) {{
            startersEl.style.display = 'none';
        }}

        var msgDiv = document.createElement('div');
        msgDiv.className = 'bov-chat-msg ' + role;
        msgDiv.innerHTML = html;
        messagesEl.appendChild(msgDiv);

        // Sources accordion (for bot messages)
        if (role === 'bot' && sources && sources.length > 0) {{
            var srcDiv = document.createElement('div');
            srcDiv.className = 'bov-chat-sources';
            var uniqueSources = [];
            var seen = {{}};
            sources.forEach(function(s) {{
                var key = s.source + ' - ' + s.page;
                if (!seen[key]) {{ seen[key] = true; uniqueSources.push(key); }}
            }});
            srcDiv.innerHTML = '<span>&#128196; Sources (' + uniqueSources.length + ')</span>' +
                '<div class="bov-chat-sources-list">' + uniqueSources.join('<br>') + '</div>';
            srcDiv.addEventListener('click', function() {{ srcDiv.classList.toggle('expanded'); }});
            messagesEl.appendChild(srcDiv);
        }}

        // Disclaimer (for bot messages)
        if (role === 'bot') {{
            var discDiv = document.createElement('div');
            discDiv.className = 'bov-chat-disclaimer';
            discDiv.textContent = DISCLAIMER;
            messagesEl.appendChild(discDiv);
        }}

        messagesEl.scrollTop = messagesEl.scrollHeight;
    }}

    // --- Typing indicator ---
    function showTyping() {{
        var el = document.createElement('div');
        el.className = 'bov-chat-typing';
        el.id = 'bovTyping';
        el.innerHTML = '<span></span><span></span><span></span>';
        messagesEl.appendChild(el);
        messagesEl.scrollTop = messagesEl.scrollHeight;
    }}
    function hideTyping() {{
        var el = document.getElementById('bovTyping');
        if (el) el.remove();
    }}

    // --- Show rate limit ---
    function showRateLimit() {{
        inputArea.style.display = 'none';
        var limitDiv = document.createElement('div');
        limitDiv.className = 'bov-chat-limit-msg';
        limitDiv.innerHTML = 'You have reached the question limit for this session. ' +
            'For more information, <a href="#contact">contact the LAAA Team</a>.';
        widget.appendChild(limitDiv);
    }}

    // --- Send message ---
    function sendMessage(text) {{
        text = (text || '').trim();
        if (!text || isStreaming) return;
        if (messageCount >= MAX_MESSAGES) {{ showRateLimit(); return; }}

        inputEl.value = '';
        addMessage('user', text.replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;'));
        messageCount++;

        // Check for pre-computed answer
        if (PRECOMPUTED[text]) {{
            var pc = PRECOMPUTED[text];
            history.push({{ role: 'user', content: text }});
            history.push({{ role: 'assistant', content: pc.answer }});
            addMessage('bot', renderMarkdown(pc.answer), pc.sources || []);
            return;
        }}

        // Call the worker
        isStreaming = true;
        sendBtn.disabled = true;
        showTyping();

        // Truncate history to last 6 turns (12 messages) to control token usage
        var trimmedHistory = history.slice(-12);

        fetch(WORKER_URL + '/chat', {{
            method: 'POST',
            headers: {{ 'Content-Type': 'application/json' }},
            body: JSON.stringify({{
                question: text,
                namespace: NAMESPACE,
                history: trimmedHistory
            }})
        }})
        .then(function(response) {{
            if (!response.ok) {{
                throw new Error('Server error: ' + response.status);
            }}
            hideTyping();

            var reader = response.body.getReader();
            var decoder = new TextDecoder();
            var fullText = '';
            var sources = [];
            var botDiv = document.createElement('div');
            botDiv.className = 'bov-chat-msg bot';
            messagesEl.appendChild(botDiv);

            function readStream() {{
                return reader.read().then(function(result) {{
                    if (result.done) {{
                        // Finalize
                        botDiv.innerHTML = renderMarkdown(fullText);
                        messagesEl.scrollTop = messagesEl.scrollHeight;
                        history.push({{ role: 'user', content: text }});
                        history.push({{ role: 'assistant', content: fullText }});
                        isStreaming = false;
                        sendBtn.disabled = false;

                        // Add sources + disclaimer
                        if (sources.length > 0) {{
                            var srcDiv = document.createElement('div');
                            srcDiv.className = 'bov-chat-sources';
                            var uniqueSources = [];
                            var seen = {{}};
                            sources.forEach(function(s) {{
                                var key = s.source + (s.page ? ' - ' + s.page : '');
                                if (!seen[key]) {{ seen[key] = true; uniqueSources.push(key); }}
                            }});
                            srcDiv.innerHTML = '<span>&#128196; Sources (' + uniqueSources.length + ')</span>' +
                                '<div class="bov-chat-sources-list">' + uniqueSources.join('<br>') + '</div>';
                            srcDiv.addEventListener('click', function() {{ srcDiv.classList.toggle('expanded'); }});
                            messagesEl.appendChild(srcDiv);
                        }}
                        var discDiv = document.createElement('div');
                        discDiv.className = 'bov-chat-disclaimer';
                        discDiv.textContent = DISCLAIMER;
                        messagesEl.appendChild(discDiv);
                        return;
                    }}

                    var chunk = decoder.decode(result.value, {{ stream: true }});
                    // Parse SSE lines
                    var lines = chunk.split('\\n');
                    lines.forEach(function(line) {{
                        line = line.trim();
                        if (line.startsWith('data: ')) {{
                            var data = line.substring(6);
                            if (data === '[DONE]') return;
                            try {{
                                var parsed = JSON.parse(data);
                                if (parsed.type === 'text') {{
                                    fullText += parsed.content;
                                    botDiv.innerHTML = renderMarkdown(fullText) + '<span style="opacity:0.4;">&#9608;</span>';
                                    messagesEl.scrollTop = messagesEl.scrollHeight;
                                }} else if (parsed.type === 'sources') {{
                                    sources = parsed.sources || [];
                                }} else if (parsed.type === 'error') {{
                                    fullText = 'Sorry, I encountered an error. Please try again.';
                                    botDiv.innerHTML = renderMarkdown(fullText);
                                }}
                            }} catch(e) {{}}
                        }}
                    }});

                    return readStream();
                }});
            }}

            return readStream();
        }})
        .catch(function(err) {{
            hideTyping();
            addMessage('bot', 'Sorry, I could not reach the server. Please try again in a moment.');
            isStreaming = false;
            sendBtn.disabled = false;
            console.error('BOV Chat error:', err);
        }});
    }}
}})();
</script>
'''


# ============================================================
# BUILD INTEGRATION HELPER
# ============================================================

def run_rag_pipeline(
    docs_dir: str,
    namespace: str,
    build_data: dict = None,
    verbose: bool = True
) -> tuple[list[Chunk], list[dict]]:
    """
    Run the full RAG pipeline: parse -> chunk -> embed -> upload.
    Returns (chunks, vectors) for optional post-processing (e.g., pre-computing starter answers).

    Args:
        docs_dir: Path to docs/ folder
        namespace: Pinecone namespace for this BOV
        build_data: Structured data from build script Python variables
        verbose: Print progress messages
    """
    if verbose:
        print("\n" + "=" * 60)
        print("BUILDING RAG KNOWLEDGE BASE")
        print("=" * 60)

    # 1. Parse all documents
    if verbose:
        print("\nStep 1: Parsing documents...")
    documents = parse_all_documents(docs_dir, build_data)
    if verbose:
        print(f"  Total: {len(documents)} document segments parsed")

    if not documents:
        print("  WARNING: No documents found. Chatbot will have no knowledge base.")
        return [], []

    # 2. Chunk
    if verbose:
        print("\nStep 2: Chunking documents...")
    chunks = smart_chunk(documents)
    if verbose:
        print(f"  Total: {len(chunks)} chunks created")
        tabular = sum(1 for c in chunks if c.is_tabular)
        narrative = len(chunks) - tabular
        print(f"  Tabular chunks: {tabular}, Narrative chunks: {narrative}")

    # 3. Embed
    if verbose:
        print("\nStep 3: Generating embeddings (Voyage AI)...")
    vectors = embed_chunks(chunks, namespace)
    if verbose:
        print(f"  Total: {len(vectors)} embeddings generated")

    # 4. Upload
    if verbose:
        print("\nStep 4: Uploading to Pinecone...")
    upload_vectors(vectors, namespace)
    if verbose:
        print(f"  Upload complete! Namespace: {namespace}")
        print("=" * 60 + "\n")

    return chunks, vectors
